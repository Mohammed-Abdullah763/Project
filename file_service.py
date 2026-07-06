"""
app/services/file_service.py
────────────────────────────
Handles file uploads and text extraction for:
  • PDF   (.pdf)   — via PyMuPDF (fitz)
  • DOCX  (.docx)  — via python-docx
  • TXT   (.txt)   — plain read
  • MD    (.md)    — plain read
All extracted text is cleaned and returned for AI processing.
"""

from __future__ import annotations

import io
import os
import re
import uuid
from datetime import datetime
from pathlib import Path
from typing import BinaryIO

import aiofiles
from fastapi import HTTPException, UploadFile, status

from app.core.config import settings

# ── allowed types ─────────────────────────────────────────────────────────────
ALLOWED_MIME = {
    "application/pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "text/plain",
    "text/markdown",
}
ALLOWED_EXT = {".pdf", ".docx", ".txt", ".md"}
MAX_FILE_SIZE_MB = 10
MAX_FILE_BYTES = MAX_FILE_SIZE_MB * 1024 * 1024


# ═════════════════════════════════════════════════════════════════════════════
# VALIDATION
# ═════════════════════════════════════════════════════════════════════════════
def validate_file(file: UploadFile) -> None:
    """Raise HTTPException if the file is invalid."""
    ext = Path(file.filename or "").suffix.lower()
    if ext not in ALLOWED_EXT:
        raise HTTPException(
            status_code=status.HTTP_415_UNSUPPORTED_MEDIA_TYPE,
            detail=f"File type '{ext}' is not supported. "
                   f"Allowed: {', '.join(ALLOWED_EXT)}",
        )


# ═════════════════════════════════════════════════════════════════════════════
# STORAGE
# ═════════════════════════════════════════════════════════════════════════════
async def save_upload(file: UploadFile, user_id: int) -> dict:
    """
    Save an uploaded file to disk and return metadata.

    Returns:
        {
          "file_id": str,
          "original_name": str,
          "saved_path": str,
          "size_bytes": int,
          "extension": str,
          "uploaded_at": str,
        }
    """
    validate_file(file)

    content = await file.read()
    if len(content) > MAX_FILE_BYTES:
        raise HTTPException(
            status_code=status.HTTP_413_REQUEST_ENTITY_TOO_LARGE,
            detail=f"File exceeds {MAX_FILE_SIZE_MB} MB limit.",
        )

    ext = Path(file.filename or "file").suffix.lower()
    file_id = str(uuid.uuid4())
    file_name = f"{file_id}{ext}"

    upload_dir = Path(settings.UPLOAD_DIR) / str(user_id)
    upload_dir.mkdir(parents=True, exist_ok=True)
    file_path = upload_dir / file_name

    async with aiofiles.open(file_path, "wb") as f:
        await f.write(content)

    return {
        "file_id": file_id,
        "original_name": file.filename,
        "saved_path": str(file_path),
        "size_bytes": len(content),
        "extension": ext,
        "uploaded_at": datetime.utcnow().isoformat(),
    }


# ═════════════════════════════════════════════════════════════════════════════
# TEXT EXTRACTION
# ═════════════════════════════════════════════════════════════════════════════
def extract_text_from_file(file_path: str) -> str:
    """
    Extract plain text from a saved file.
    Dispatches to the correct extractor based on extension.
    """
    path = Path(file_path)
    if not path.exists():
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="File not found on server.",
        )

    ext = path.suffix.lower()
    extractors = {
        ".pdf":  _extract_pdf,
        ".docx": _extract_docx,
        ".txt":  _extract_txt,
        ".md":   _extract_txt,
    }

    extractor = extractors.get(ext)
    if not extractor:
        raise HTTPException(
            status_code=status.HTTP_415_UNSUPPORTED_MEDIA_TYPE,
            detail=f"Cannot extract text from '{ext}' files.",
        )

    raw_text = extractor(str(path))
    return _clean_text(raw_text)


def extract_text_from_bytes(content: bytes, extension: str) -> str:
    """Extract text directly from bytes (used in tests and direct uploads)."""
    ext = extension.lower()
    if ext == ".pdf":
        return _clean_text(_extract_pdf_bytes(content))
    elif ext == ".docx":
        return _clean_text(_extract_docx_bytes(content))
    elif ext in {".txt", ".md"}:
        return _clean_text(content.decode("utf-8", errors="replace"))
    else:
        raise ValueError(f"Unsupported extension: {ext}")


# ─── PDF ──────────────────────────────────────────────────────────────────────
def _extract_pdf(file_path: str) -> str:
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(file_path)
        pages = [page.get_text("text") for page in doc]
        doc.close()
        return "\n".join(pages)
    except ImportError:
        raise HTTPException(
            status_code=500,
            detail="PyMuPDF is not installed. Run: pip install PyMuPDF",
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"PDF extraction failed: {e}")


def _extract_pdf_bytes(content: bytes) -> str:
    try:
        import fitz
        doc = fitz.open(stream=content, filetype="pdf")
        pages = [page.get_text("text") for page in doc]
        doc.close()
        return "\n".join(pages)
    except ImportError:
        raise HTTPException(status_code=500, detail="PyMuPDF not installed.")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"PDF extraction failed: {e}")


# ─── DOCX ─────────────────────────────────────────────────────────────────────
def _extract_docx(file_path: str) -> str:
    try:
        from docx import Document
        doc = Document(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also extract table content
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text)
        return "\n".join(paragraphs)
    except ImportError:
        raise HTTPException(
            status_code=500,
            detail="python-docx is not installed. Run: pip install python-docx",
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"DOCX extraction failed: {e}")


def _extract_docx_bytes(content: bytes) -> str:
    try:
        from docx import Document
        doc = Document(io.BytesIO(content))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        return "\n".join(paragraphs)
    except ImportError:
        raise HTTPException(status_code=500, detail="python-docx not installed.")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"DOCX extraction failed: {e}")


# ─── TXT / MD ─────────────────────────────────────────────────────────────────
def _extract_txt(file_path: str) -> str:
    with open(file_path, "r", encoding="utf-8", errors="replace") as f:
        return f.read()


# ─── CLEAN ────────────────────────────────────────────────────────────────────
def _clean_text(text: str) -> str:
    """
    Clean extracted text:
    - Remove excessive whitespace
    - Normalize line endings
    - Strip null bytes
    """
    text = text.replace("\x00", "")                    # null bytes
    text = re.sub(r"\r\n|\r", "\n", text)              # normalize line endings
    text = re.sub(r"[ \t]+", " ", text)                # collapse spaces/tabs
    text = re.sub(r"\n{4,}", "\n\n\n", text)           # max 3 blank lines
    lines = [line.strip() for line in text.split("\n")]
    return "\n".join(lines).strip()


# ═════════════════════════════════════════════════════════════════════════════
# METADATA HELPERS
# ═════════════════════════════════════════════════════════════════════════════
def get_file_info(file_path: str) -> dict:
    """Return basic metadata about a stored file."""
    path = Path(file_path)
    if not path.exists():
        return {}
    stat = path.stat()
    return {
        "name": path.name,
        "size_bytes": stat.st_size,
        "size_mb": round(stat.st_size / (1024 * 1024), 2),
        "extension": path.suffix.lower(),
        "modified_at": datetime.fromtimestamp(stat.st_mtime).isoformat(),
    }


def delete_file(file_path: str) -> bool:
    """Delete a file from disk. Returns True if deleted, False if not found."""
    path = Path(file_path)
    if path.exists():
        path.unlink()
        return True
    return False

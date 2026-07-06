"""
brain.py — StudyMind Pro Intelligence Layer
============================================
The "own brain" of the app:
  - Text file reading (TXT, MD, PDF, DOCX, CSV, JSON)
  - Image understanding (photos, handwriting, diagrams)
  - OCR for bad/handwritten text
  - Document summarisation and indexing
  - Knowledge base management
"""

import base64
import io
import json
import re
from pathlib import Path
from typing import Any

import config as cfg
import connections as conn

# ══════════════════════════════════════════════════════
# FILE TEXT EXTRACTION
# ══════════════════════════════════════════════════════

def extract_text_from_file(path: Path, ext: str) -> str:
    """
    Extract plain text from any supported file type.
    Handles: TXT, MD, PDF, DOCX, CSV, JSON
    """
    ext = ext.lower()

    # ── Plain text ─────────────────────────────────────
    if ext in (".txt", ".md"):
        return path.read_text(encoding="utf-8", errors="ignore")

    # ── PDF ────────────────────────────────────────────
    if ext == ".pdf":
        try:
            import PyPDF2
            with open(path, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                pages = []
                for i, page in enumerate(reader.pages):
                    text = page.extract_text() or ""
                    if text.strip():
                        pages.append(f"[Page {i+1}]\n{text}")
                return "\n\n".join(pages)
        except ImportError:
            return "[PDF support: install PyPDF2 with: pip install PyPDF2]"
        except Exception as e:
            return f"[PDF read error: {e}]"

    # ── Word document ──────────────────────────────────
    if ext == ".docx":
        try:
            from docx import Document
            doc = Document(str(path))
            parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    parts.append(para.text)
            # Also extract tables
            for table in doc.tables:
                for row in table.rows:
                    parts.append(" | ".join(cell.text for cell in row.cells))
            return "\n".join(parts)
        except ImportError:
            return "[DOCX support: install python-docx with: pip install python-docx]"
        except Exception as e:
            return f"[DOCX read error: {e}]"

    # ── CSV ────────────────────────────────────────────
    if ext == ".csv":
        try:
            import csv
            with open(path, encoding="utf-8", errors="ignore") as f:
                reader = csv.reader(f)
                rows = list(reader)
                if not rows:
                    return ""
                # Header + first 50 rows
                lines = [" | ".join(rows[0])]
                lines.append("-" * 40)
                for row in rows[1:51]:
                    lines.append(" | ".join(row))
                if len(rows) > 51:
                    lines.append(f"... and {len(rows)-51} more rows")
                return "\n".join(lines)
        except Exception as e:
            return f"[CSV read error: {e}]"

    # ── JSON ───────────────────────────────────────────
    if ext == ".json":
        try:
            with open(path, encoding="utf-8") as f:
                data = json.load(f)
            return json.dumps(data, indent=2)[:5000]
        except Exception as e:
            return f"[JSON read error: {e}]"

    return "[Unsupported file type]"

# ══════════════════════════════════════════════════════
# IMAGE ANALYSIS
# ══════════════════════════════════════════════════════

async def analyse_image(image_bytes: bytes, media_type: str, task: str = "general") -> dict:
    """
    Analyse an image using Claude Vision.

    Tasks:
      general   — describe what's in the image
      ocr       — extract all text (handles bad/handwriting)
      diagram   — explain diagrams, charts, mind maps
      notes     — read study notes / whiteboard photos
      math      — solve equations visible in image
    """
    prompts = {
        "general": (
            "Describe this image in detail. If there is any text, transcribe it exactly. "
            "Explain any diagrams, charts, or visual elements. Be thorough."
        ),
        "ocr": (
            "Your task is OCR — extract ALL text visible in this image, exactly as written. "
            "Include handwritten text, printed text, labels, captions, and annotations. "
            "If handwriting is unclear, make your best attempt and mark uncertain parts with [?]. "
            "Preserve the original structure and layout as much as possible."
        ),
        "diagram": (
            "This image contains a diagram, chart, or visual. "
            "1. Identify what type it is (flowchart, graph, mind map, etc.) "
            "2. Explain what it shows and its key components "
            "3. Describe relationships and connections "
            "4. State the main insight or conclusion it conveys"
        ),
        "notes": (
            "These are study notes or a whiteboard. "
            "1. Transcribe all text, preserving headings and structure "
            "2. Identify the main topic and subtopics "
            "3. Note any diagrams, formulas, or special notation "
            "4. Summarise the key learning points"
        ),
        "math": (
            "Extract and solve any mathematical content in this image. "
            "1. Transcribe equations exactly as written "
            "2. Identify the type of problem "
            "3. Show the solution step by step "
            "4. State the final answer clearly"
        ),
    }

    prompt = prompts.get(task, prompts["general"])

    try:
        result = await conn.ai_vision(image_bytes, media_type, prompt)
        return {
            "success": True,
            "task": task,
            "result": result,
            "media_type": media_type,
        }
    except Exception as e:
        return {
            "success": False,
            "task": task,
            "result": f"Image analysis failed: {str(e)}",
            "error": str(e),
        }

# ══════════════════════════════════════════════════════
# OCR (Tesseract fallback if Vision API unavailable)
# ══════════════════════════════════════════════════════

def ocr_fallback(image_bytes: bytes) -> str:
    """
    Tesseract OCR as fallback when Vision API is unavailable.
    Install: pip install pytesseract pillow
    System:  sudo apt-get install tesseract-ocr
    """
    if not cfg.OCR_ENABLED:
        return ""
    try:
        import pytesseract
        from PIL import Image
        img = Image.open(io.BytesIO(image_bytes))
        return pytesseract.image_to_string(img, lang=cfg.OCR_LANGUAGE)
    except ImportError:
        return "[OCR: install pytesseract and pillow]"
    except Exception as e:
        return f"[OCR error: {e}]"

# ══════════════════════════════════════════════════════
# AI FEATURE PROMPTS
# ══════════════════════════════════════════════════════

async def summarize(text: str) -> dict:
    raw = await conn.ai_text(
        "You are an expert academic summarizer. Return ONLY valid JSON, no markdown fences.",
        f'Summarize this study material. Return JSON:\n'
        f'{{"tldr":"one sentence","key_points":["p1","p2","p3","p4","p5"],'
        f'"concepts":["c1","c2","c3"],"importance":"why this matters in 2 sentences"}}\n\n'
        f'Material:\n{text[:4500]}'
    )
    return conn._parse_json(raw)

async def flashcards(text: str) -> list:
    raw = await conn.ai_text(
        "You are a flashcard expert. Return ONLY valid JSON array, no markdown.",
        f'Create 10 high-quality flashcards. Return JSON:\n'
        f'[{{"q":"question","a":"answer","hint":"short hint","difficulty":"easy|medium|hard"}}]\n\n'
        f'Material:\n{text[:4500]}'
    )
    return conn._parse_json(raw)

async def quiz(text: str) -> list:
    raw = await conn.ai_text(
        "You are a quiz master. Return ONLY valid JSON array, no markdown.",
        f'Create 8 multiple-choice questions. Return JSON:\n'
        f'[{{"q":"question","options":["A","B","C","D"],"answer":0,"explanation":"why correct"}}]\n'
        f'(answer is 0-indexed)\n\nMaterial:\n{text[:4500]}'
    )
    return conn._parse_json(raw)

async def mindmap(text: str) -> dict:
    raw = await conn.ai_text(
        "You are a knowledge-structure expert. Return ONLY valid JSON, no markdown.",
        f'Build a mind map. Return JSON:\n'
        f'{{"root":"main topic","branches":[{{"name":"branch","leaves":["l1","l2","l3"]}}]}}\n'
        f'4-6 branches, 3-5 leaves.\n\nMaterial:\n{text[:4500]}'
    )
    return conn._parse_json(raw)

async def key_terms(text: str) -> list:
    raw = await conn.ai_text(
        "You are a glossary expert. Return ONLY valid JSON array, no markdown.",
        f'Extract 12 key terms. Return JSON:\n'
        f'[{{"term":"name","definition":"1-2 sentence definition"}}]\n\nMaterial:\n{text[:4500]}'
    )
    return conn._parse_json(raw)

async def study_plan(text: str) -> list:
    raw = await conn.ai_text(
        "You are a learning designer. Return ONLY valid JSON array, no markdown.",
        f'Create a 7-day study plan. Return JSON:\n'
        f'[{{"day":1,"title":"Day title","tasks":"activities","goal":"what student will know"}}]\n\n'
        f'Material:\n{text[:4500]}'
    )
    return conn._parse_json(raw)
